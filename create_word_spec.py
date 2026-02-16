#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
デジタルサイネージシステム 要件定義書（Word形式）
クライアント提出用
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """セルの背景色を設定"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(doc):
    """水平線を追加"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('─' * 80)
    run.font.color.rgb = RGBColor(200, 200, 200)
    run.font.size = Pt(8)

def create_spec_document():
    doc = Document()
    
    # フォント設定
    style = doc.styles['Normal']
    style.font.name = 'メイリオ'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'メイリオ')
    
    # ===== 表紙 =====
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('デジタルサイネージシステム')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(26, 54, 93)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('要件定義書')
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(49, 130, 206)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # バージョン情報
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('Ver 1.0')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 会社情報
    company = doc.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = company.add_run('2026年2月')
    run.font.size = Pt(14)
    
    company2 = doc.add_paragraph()
    company2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = company2.add_run('千歳開発')
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_page_break()
    
    # ===== 目次 =====
    h = doc.add_heading('目次', level=1)
    h.runs[0].font.color.rgb = RGBColor(26, 54, 93)
    
    toc_items = [
        '1. プロジェクト概要',
        '2. システム構成',
        '3. 画面仕様',
        '   3.1 行動予定表',
        '   3.2 施工サイクル表',
        '   3.3 協力業者作業一覧',
        '   3.4 月間行事予定表',
        '   3.5 ポスター表示',
        '   3.6 管理画面',
        '4. 表示サイクル仕様',
        '5. 非機能要件',
        '6. 開発スケジュール',
        '7. 納品物',
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.line_spacing = 1.8
    
    doc.add_page_break()
    
    # ===== 1. プロジェクト概要 =====
    h = doc.add_heading('1. プロジェクト概要', level=1)
    h.runs[0].font.color.rgb = RGBColor(26, 54, 93)
    
    h2 = doc.add_heading('1.1 目的', level=2)
    h2.runs[0].font.color.rgb = RGBColor(49, 130, 206)
    
    doc.add_paragraph(
        '現場事務所内に設置されたモニターで、行動予定・協力業者・月間予定・施工サイクル等の'
        '情報をデジタル表示するサイネージシステムを構築します。'
    )
    
    h2 = doc.add_heading('1.2 システム特性', level=2)
    h2.runs[0].font.color.rgb = RGBColor(49, 130, 206)
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    data = [
        ('項目', '内容'),
        ('ネットワーク', '完全オンプレミス環境（インターネット接続なし）'),
        ('動作環境', '単一Windows PC上で動作'),
        ('外部通信', '外部通信を行わない閉じた構成'),
    ]
    
    for i, (col1, col2) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = col1
        row.cells[1].text = col2
        if i == 0:
            set_cell_shading(row.cells[0], '1a365d')
            set_cell_shading(row.cells[1], '1a365d')
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[1].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    h2 = doc.add_heading('1.3 制約事項', level=2)
    h2.runs[0].font.color.rgb = RGBColor(49, 130, 206)
    
    constraints = [
        '天気予報等の外部API連携は不可',
        '縦置きモニターと横置きモニターの表を同一サイクルに混在させることは不可',
    ]
    for c in constraints:
        p = doc.add_paragraph(c, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 2. システム構成 =====
    h = doc.add_heading('2. システム構成', level=1)
    h.runs[0].font.color.rgb = RGBColor(26, 54, 93)
    
    h2 = doc.add_heading('2.1 技術スタック', level=2)
    h2.runs[0].font.color.rgb = RGBColor(49, 130, 206)
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    
    data = [
        ('項目', '技術', '備考'),
        ('フレームワーク', 'Electron', 'デスクトップアプリとして動作'),
        ('フロントエンド', 'React または Vue.js', 'コンポーネント指向'),
        ('データベース', 'SQLite', 'ファイルベース、軽量'),
        ('設定管理', 'YAML または JSON', '画面定義の外部化'),
    ]
    
    for i, (col1, col2, col3) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = col1
        row.cells[1].text = col2
        row.cells[2].text = col3
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, '1a365d')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    h2 = doc.add_heading('2.2 動作環境', level=2)
    h2.runs[0].font.color.rgb = RGBColor(49, 130, 206)
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    data = [
        ('項目', '要件'),
        ('OS', 'Windows 10/11'),
        ('メモリ', '4GB以上推奨'),
        ('ストレージ', '1GB以上の空き容量'),
        ('ディスプレイ', '縦置き: 1080×1920 / 横置き: 1920×1080'),
    ]
    
    for i, (col1, col2) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = col1
        row.cells[1].text = col2
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, '1a365d')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ===== 3. 画面仕様 =====
    h = doc.add_heading('3. 画面仕様', level=1)
    h.runs[0].font.color.rgb = RGBColor(26, 54, 93)
    
    # 画面一覧表
    h2 = doc.add_heading('画面一覧', level=2)
    h2.runs[0].font.color.rgb = RGBColor(49, 130, 206)
    
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    
    data = [
        ('No', '画面名', '向き', '用途'),
        ('1', '行動予定表', '縦置き', '社員の行動状況表示'),
        ('2', '施工サイクル表', '縦置き', '当日の作業工程・ステータス表示'),
        ('3', '協力業者作業一覧', '横置き', '当日入場の協力業者表示'),
        ('4', '月間行事予定表', '横置き', '月間カレンダー形式の予定表示'),
        ('5', 'ポスター表示', '両対応', '社内ポスター・掲示物のスライドショー'),
        ('6', '管理画面', '-', 'データ登録・設定変更用'),
    ]
    
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, text in enumerate(row_data):
            row.cells[j].text = text
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, '1a365d')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ----- 3.1 行動予定表 -----
    h2 = doc.add_heading('3.1 行動予定表', level=2)
    h2.runs[0].font.color.rgb = RGBColor(49, 130, 206)
    
    p = doc.add_paragraph()
    run = p.add_run('【縦置き 75インチ (1080×1920)】')
    run.font.bold = True
    run.font.color.rgb = RGBColor(159, 122, 234)
    
    doc.add_paragraph('社員の行動状況（氏名・行先・使用車両・帰社予定時刻）を一覧表示します。')
    
    # 画像挿入
    img_path = 'docs/screenshots/01_行動予定表.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    h3 = doc.add_heading('表示項目', level=3)
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    
    data = [
        ('項目', '型', '必須', '説明'),
        ('氏名', '文字列', '○', '社員名'),
        ('行先', '文字列', '○', '外出先・「在席」等'),
        ('使用車両', '文字列', '-', '社用車のナンバー等'),
        ('帰社予定', '時刻', '-', 'HH:MM形式'),
    ]
    
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, text in enumerate(row_data):
            row.cells[j].text = text
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, '2d3748')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    h3 = doc.add_heading('動的スケーリング仕様', level=3)
    doc.add_paragraph('表示人数に応じてフォントサイズを自動調整します。')
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    data = [
        ('行数', 'フォントサイズ'),
        ('10行以下', '26px（基本）'),
        ('11〜15行', '22px'),
        ('16〜20行', '18px'),
        ('21行以上', '16px（スクロール検討）'),
    ]
    for i, (col1, col2) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = col1
        row.cells[1].text = col2
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, '2d3748')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ----- 3.2 施工サイクル表 -----
    h2 = doc.add_heading('3.2 施工サイクル表', level=2)
    h2.runs[0].font.color.rgb = RGBColor(49, 130, 206)
    
    p = doc.add_paragraph()
    run = p.add_run('【縦置き】')
    run.font.bold = True
    run.font.color.rgb = RGBColor(159, 122, 234)
    
    doc.add_paragraph('当日の作業内容とステータス（完了/作業中/未着手/要確認）を表形式で表示します。')
    
    img_path = 'docs/screenshots/02_施工サイクル表.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    h3 = doc.add_heading('ステータス表示', level=3)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    data = [
        ('ステータス', '記号', '色'),
        ('完了', '✓', '緑 (#38a169)'),
        ('作業中', '◐', '黄 (#d69e2e)'),
        ('未着手', '○', 'グレー (#a0aec0)'),
        ('要確認', '!', '赤 (#e53e3e)'),
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, text in enumerate(row_data):
            row.cells[j].text = text
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, '2d3748')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ----- 3.3 協力業者作業一覧 -----
    h2 = doc.add_heading('3.3 協力業者作業一覧', level=2)
    h2.runs[0].font.color.rgb = RGBColor(49, 130, 206)
    
    p = doc.add_paragraph()
    run = p.add_run('【横置き】')
    run.font.bold = True
    run.font.color.rgb = RGBColor(66, 153, 225)
    
    doc.add_paragraph('当日入場の協力業者名・作業内容・人数をカード形式で表示します。4列グリッドで配置し、業者数に応じてカードサイズを自動調整します。')
    
    img_path = 'docs/screenshots/03_協力業者一覧.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    h3 = doc.add_heading('表示項目', level=3)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    data = [
        ('項目', '型', '必須'),
        ('業者名', '文字列', '○'),
        ('作業内容', '文字列', '-'),
        ('人数', '数値', '-'),
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, text in enumerate(row_data):
            row.cells[j].text = text
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, '2d3748')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('※画面上部に「入場業者数」「入場人数」のサマリーを表示します。')
    
    doc.add_page_break()
    
    # ----- 3.4 月間行事予定表 -----
    h2 = doc.add_heading('3.4 月間行事予定表', level=2)
    h2.runs[0].font.color.rgb = RGBColor(49, 130, 206)
    
    p = doc.add_paragraph()
    run = p.add_run('【横置き 75インチ (1920×1080)】')
    run.font.bold = True
    run.font.color.rgb = RGBColor(66, 153, 225)
    
    doc.add_paragraph('日付・曜日・行事予定をカレンダー形式で表示します。当日の行はハイライト表示されます。')
    
    img_path = 'docs/screenshots/04_月間行事予定表.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    h3 = doc.add_heading('イベントタグ', level=3)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    data = [
        ('タグ', '用途', '背景色'),
        ('重要', '重要イベント', '#fed7d7（薄い赤）'),
        ('定例', '定例会議等', '#c6f6d5（薄い緑）'),
        ('作業', '作業予定', '#bee3f8（薄い青）'),
        ('休日', '休日・祝日', '#feebc8（薄いオレンジ）'),
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, text in enumerate(row_data):
            row.cells[j].text = text
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, '2d3748')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ----- 3.5 ポスター表示 -----
    h2 = doc.add_heading('3.5 ポスター表示', level=2)
    h2.runs[0].font.color.rgb = RGBColor(49, 130, 206)
    
    p = doc.add_paragraph()
    run = p.add_run('【縦置き・横置き両対応】')
    run.font.bold = True
    run.font.color.rgb = RGBColor(72, 187, 120)
    
    doc.add_paragraph('社内ポスター・掲示物をスライドショー形式で表示します。')
    
    doc.add_paragraph()
    h3 = doc.add_heading('対応フォーマット', level=3)
    formats = ['JPEG (.jpg, .jpeg)', 'PNG (.png)', 'PDF (.pdf) ※1ファイル=1ページ前提']
    for f in formats:
        doc.add_paragraph(f, style='List Bullet')
    
    doc.add_paragraph()
    h3 = doc.add_heading('表示パターン（選択式）', level=3)
    
    # パターンA
    p = doc.add_paragraph()
    run = p.add_run('■ パターンA：フルスクリーン切替方式【推奨】')
    run.font.bold = True
    run.font.color.rgb = RGBColor(56, 161, 105)
    
    doc.add_paragraph('表とポスターを交互に切り替えて表示します。')
    doc.add_paragraph('動作: 表 → 表 → ポスター → 表 → ポスター → ループ')
    
    img_path = 'docs/screenshots/05_パターンA.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('✅ メリット: ')
    run.font.color.rgb = RGBColor(56, 161, 105)
    p.add_run('ポスターが大きく見やすい / 実装がシンプル')
    
    p = doc.add_paragraph()
    run = p.add_run('❌ デメリット: ')
    run.font.color.rgb = RGBColor(229, 62, 62)
    p.add_run('表とポスターを同時に見られない')
    
    doc.add_paragraph()
    
    # パターンB
    p = doc.add_paragraph()
    run = p.add_run('■ パターンB：分割表示方式')
    run.font.bold = True
    run.font.color.rgb = RGBColor(113, 128, 150)
    
    doc.add_paragraph('画面を上下に分割し、表とポスターを同時表示します。（上70%: 表 / 下30%: ポスター）')
    
    img_path = 'docs/screenshots/06_パターンB.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('✅ メリット: ')
    run.font.color.rgb = RGBColor(56, 161, 105)
    p.add_run('表とポスターを同時に確認できる')
    
    p = doc.add_paragraph()
    run = p.add_run('❌ デメリット: ')
    run.font.color.rgb = RGBColor(229, 62, 62)
    p.add_run('表示領域が狭くなる / レイアウト調整が複雑')
    
    doc.add_page_break()
    
    # ----- 3.6 管理画面 -----
    h2 = doc.add_heading('3.6 管理画面', level=2)
    h2.runs[0].font.color.rgb = RGBColor(49, 130, 206)
    
    doc.add_paragraph('データ登録・設定変更を行う管理画面を提供します。サイネージ端末からのみアクセス可能です。')
    
    doc.add_paragraph()
    h3 = doc.add_heading('機能一覧', level=3)
    
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    data = [
        ('機能', '説明'),
        ('行動予定表管理', '社員の追加・編集・削除、行動状況の更新'),
        ('施工サイクル管理', '作業項目の追加・編集・削除・並び替え、ステータス更新'),
        ('協力業者管理', '業者の追加・編集・削除・並び替え'),
        ('月間予定管理', '年月選択、日付ごとの予定入力・編集・削除'),
        ('ポスター管理', 'ディレクトリ指定、プレビュー、表示順設定'),
        ('表示設定', '切り替えサイクル秒数、各画面の表示/非表示、タッチモード切替'),
        ('画面定義管理', '新規画面の追加（将来拡張用）'),
    ]
    for i, (col1, col2) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = col1
        row.cells[1].text = col2
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, '1a365d')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ===== 4. 表示サイクル仕様 =====
    h = doc.add_heading('4. 表示サイクル仕様', level=1)
    h.runs[0].font.color.rgb = RGBColor(26, 54, 93)
    
    doc.add_paragraph('各モニターで表示する画面を順番に切り替えます。切り替え間隔は設定画面から変更可能です。')
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    data = [
        ('モニター', '切替間隔', '表示順序'),
        ('縦置き', '10秒', '行動予定表 → 施工サイクル → ポスター → ループ'),
        ('横置き', '10秒', '協力業者 → 月間予定 → ポスター → ループ'),
        ('ポスター間', '5秒', 'ポスター画像を順番に切り替え'),
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, text in enumerate(row_data):
            row.cells[j].text = text
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, '1a365d')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('※注意: ')
    run.font.bold = True
    p.add_run('縦置き画面と横置き画面は別サイクルとして管理し、同一サイクル内での縦横混在はできません。')
    
    doc.add_page_break()
    
    # ===== 5. 非機能要件 =====
    h = doc.add_heading('5. 非機能要件', level=1)
    h.runs[0].font.color.rgb = RGBColor(26, 54, 93)
    
    h2 = doc.add_heading('5.1 パフォーマンス', level=2)
    h2.runs[0].font.color.rgb = RGBColor(49, 130, 206)
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    data = [
        ('項目', '要件'),
        ('画面切り替え', '500ms以内'),
        ('起動時間', '10秒以内'),
        ('メモリ使用量', '500MB以下'),
    ]
    for i, (col1, col2) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = col1
        row.cells[1].text = col2
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, '2d3748')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    h2 = doc.add_heading('5.2 信頼性', level=2)
    h2.runs[0].font.color.rgb = RGBColor(49, 130, 206)
    
    items = [
        '24時間365日連続稼働を想定',
        'エラー発生時は自動復帰（クラッシュ時は自動再起動）',
        'データベースの定期バックアップ機能',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    
    h2 = doc.add_heading('5.3 セキュリティ', level=2)
    h2.runs[0].font.color.rgb = RGBColor(49, 130, 206)
    
    items = [
        'インターネット非接続による物理的セキュリティ確保',
        'ローカルネットワークからのアクセスも原則不可',
        'ファイルシステムへのアクセスは最小限に制限',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 6. 開発スケジュール =====
    h = doc.add_heading('6. 開発スケジュール', level=1)
    h.runs[0].font.color.rgb = RGBColor(26, 54, 93)
    
    doc.add_paragraph('想定開発期間: 約8週間')
    doc.add_paragraph()
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    data = [
        ('フェーズ', '期間', '内容'),
        ('Phase 1', '2週間', '基盤構築（Electron + DB + 設定読み込み）'),
        ('Phase 2', '2週間', 'サイネージ画面5種類の実装'),
        ('Phase 3', '2週間', '管理画面の実装'),
        ('Phase 4', '1週間', '結合テスト・調整'),
        ('Phase 5', '1週間', '納品・導入サポート'),
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, text in enumerate(row_data):
            row.cells[j].text = text
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, '1a365d')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # ===== 7. 納品物 =====
    h = doc.add_heading('7. 納品物', level=1)
    h.runs[0].font.color.rgb = RGBColor(26, 54, 93)
    
    items = [
        '実行可能ファイル（Windows用 .exe）',
        'ソースコード一式',
        '操作マニュアル',
        '設定変更ガイド（画面追加手順含む）',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # フッター
    add_horizontal_line(doc)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('本資料に関するお問い合わせは 千歳開発 までご連絡ください。')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # 保存
    output_path = 'docs/デジタルサイネージシステム_要件定義書.docx'
    doc.save(output_path)
    print(f'✅ 保存完了: {output_path}')

if __name__ == '__main__':
    os.chdir('/Users/ogikubo/Desktop/digital-signage')
    create_spec_document()
